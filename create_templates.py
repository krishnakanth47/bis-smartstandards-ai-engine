import os
import shutil
from docx import Document

os.makedirs("frontend/assets", exist_ok=True)

# 1. Create DOCX Template
print("Generating DOCX template...")
doc = Document()
doc.add_heading('BIS Product Information Form', 0)

doc.add_heading('Product Details', level=1)
doc.add_paragraph('Product Name: ____________________________________')
doc.add_paragraph('Material Category (e.g. Cement, Steel, Concrete): _________________')

doc.add_heading('Product Description', level=1)
doc.add_paragraph('Please describe your product, its components, and manufacturing process:')
doc.add_paragraph('\n\n\n')

doc.add_heading('Intended Usage / Application', level=1)
doc.add_paragraph('Where and how will this product be used?')
doc.add_paragraph('\n\n')

doc.add_heading('Key Properties', level=1)
doc.add_paragraph('e.g. 33 Grade, Fe 415, Calcined Clay based, Corrugated')
doc.add_paragraph('\n\n')

doc.add_heading('Additional Notes', level=1)
doc.add_paragraph('\n\n')

doc.save('frontend/assets/template.docx')
print("Saved frontend/assets/template.docx")

# 2. Copy the existing PDF to the assets folder
pdf_source = "BIS_Product_Information_Form.pdf.pdf"
pdf_dest = "frontend/assets/template.pdf"

if os.path.exists(pdf_source):
    shutil.copy(pdf_source, pdf_dest)
    print(f"Copied {pdf_source} to {pdf_dest}")
else:
    print(f"Warning: {pdf_source} not found. Please place the PDF template in the root directory.")

print("Template generation complete!")
